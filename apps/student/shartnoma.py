from io import BytesIO
import os
from docx.api import Document
from docx.document import Document as Doc
from docx.table import _Cell
from docx.shared import Pt
from docx2pdf import convert
import qrcode
from docx.shared import Cm


def create_shartnoma(name, id, passport, faculty, number, date, price, mode, template):
    COMMAND = "libreoffice --headless --convert-to pdf {doc} --outdir {out}"

    # Saqlangan word fileni yoki template wordni o'qivolamiz
    word_file = open(f".{template}", "rb")

    # Ochilgan fileni Word document sifatida ochvoamiz
    doc: Doc = Document(word_file)

    # Natija uchun `On Memory` file ochib olamiz bu faqat ramda yani hotirada saqlanadi
    # open("res.docx","wb") qilib ochsaham bo'ladi. Bunda file diskga yoziladi.
    #
    res = open(f"./files/agreements/{id}.docx", "wb")
    q = qrcode.make(f'http://185.65.202.40:1009/files/agreements/{id}.pdf')
    q.save('./files/qrcode.png')
    # Word fileni ichida har bir elementni paragraphs orqali olib uni forloopga qo'yamiz
    for paragraph in doc.paragraphs:
        # paragraphni ichidigi text olib olamiz
        # PS: :str variableni typini ko'rsatish uchun berildi
        t: str = paragraph.text

        # Berilgan keywordlar paragraphni ichida mavjud yoki yo'qligini tekshiramiz
        if (
                ("{id}" in t)
                or ("{name}" in t)
                or ("{address}" in t)
                or ("{passport}" in t)
                or ("{university}" in t)
                or ("{faculty}" in t)
                or ("{number}" in t)
                or ("{end}" in t)
                or ("{delta}" in t)
                or ("{mode}" in t)
                or ("{lang}" in t)
                or ("{date}" in t)
                or ("{price}" in t)
        ):
            # Agar qaysidir keyword paragraphning ichida uchrasa uni replace qiladi.
            paragraph.text = (
                paragraph.text.replace("{id}", "4%06d" % int(f"{id}"))
                .replace("{name}", f"{name}")
                .replace("{date}", f"{date}")
                .replace("{address}", f"")
                .replace("{passport}", f"{passport}")
                .replace("{university}", f"universuty")
                .replace("{faculty}", f"{faculty}")
                .replace("{price}", f"{price}")
                .replace("{number}", f"+{number}")
                .replace("{end}", "2027")
                .replace("{mode}", f"{mode}")
                .replace("{delta}", f"{4}")
            )

        # Word Document da bazab table'lar ham bo'ladi. Ularni paragraphs bilan ololmimiz.
    # Ularni .tables bilan olinadi

    for table in doc.tables:
        # Har bir table column va rowdan tuzilgan bo'ladi ularni for loop ga qo'ymaiz
        for column in table.columns:
            for row in table.rows:
                # cell bu table ichidigi katak
                cell: _Cell = table.cell(row._index, column._index)
                # cell ni ichida paragraphlar bo'ladi
                for p in cell.paragraphs:
                    t: str = p.text
                    # Berilgan keywordlar paragraphni ichida mavjud yoki yo'qligini tekshiramiz
                    if (
                            ("{id}" in t)
                            or ("{name}" in t)
                            or ("{address}" in t)
                            or ("{passport}" in t)
                            or ("{university}" in t)
                            or ("{faculty}" in t)
                            or ("{price}" in t)
                            or ("{number}" in t)
                            or ("{end}" in t)
                            or ("{delta}" in t)
                            or ("{mode}" in t)
                            or ("{lang}" in t)
                            or ("{date}" in t)
                    ):
                        # Agar qaysidir keyword paragraphning ichida uchrasa uni replace qiladi.
                        cell.text = (
                            cell.text.replace("{id}", f"{id}")
                            .replace("{name}", f"{name}")
                            .replace("{address}", f"")
                            .replace("{passport}", f"{passport}")
                            .replace("{university}", f"universuty")
                            .replace("{faculty}", f"{faculty}")
                            .replace("{price}", f"{price}")
                            .replace("{number}", f"+{number}")
                            .replace("{end}", "2027")
                            .replace("{mode}", f"{mode}")
                            .replace("{lang}", "O'zbek")
                            .replace("{end}", "2027")
                        )

    # Iterate over paragraphs
    for paragraph in doc.paragraphs:
        if '{qr}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{qr}', "")
            run = paragraph.add_run()
            run.add_picture(f'./files/qrcode.png', width=Cm(3))

        # Access the runs within the paragraph
        for run in paragraph.runs:
            # Modify the font properties
            font = run.font
            font.name = "Tmes New Roman"
            font.size = Pt(8)

        for table in doc.tables:
            # Har bir table column va rowdan tuzilgan bo'ladi ularni for loop ga qo'ymaiz
            for column in table.columns:
                for row in table.rows:
                    # cell bu table ichidigi katak
                    cell: _Cell = table.cell(row._index, column._index)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # Modify the font properties
                            font = run.font
                            font.name = "Tmes New Roman"
                            font.size = Pt(8)
                        if '{qr}' in paragraph.text:
                            paragraph.text = paragraph.text.replace('{qr}', "")
                            run = paragraph.add_run()
                            run.add_picture(f'./files/qrcode.png', width=Cm(3))

    doc.save(res)
    # convert(f'./agreements/{id}.docx', f'./agreements/{id}.pdf')
    os.system(f"libreoffice --headless --convert-to pdf ./files/agreements/{id}.docx --outdir ./files/agreements")
